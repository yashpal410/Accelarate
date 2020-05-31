''' Program to search keywords in PDFs and save the corresponding sentence
    to Excel Worksheet.  '''
import os
import glob
from docx import Document
import re
import PyPDF2 as p2
from openpyxl import Workbook, load_workbook
from wordextracts import update_docx

folder_path = r"F:\Yashraj\yes\prac" #Location of the PDF file(s)
workbook = load_workbook(filename="words.xlsx") #Enter filename of keyword
sheet = workbook.active

wb = Workbook()
ws = wb.active

ws['A1'] = 'Document Name'
ws['B1'] = 'Keyword'
ws['C1'] = 'Page No'
ws['D1'] = 'Content'

count = 1

def excelupdate (count,filename,i,word,str1):
    ws['A' + str(count)] = filename.split("\\")[-1]   
    ws['B' + str(count)] = word                       
    ws['C' + str(count)] = i + 1                      
    ws['D' + str(count)] = str1                   

def previous(lists,str1,j):
    for k in range(j-1,0,-1):                                                               
        try:                                                                                
            if len(lists[k].split(". ")) >1:                                                
                sent = lists[k].split(". ")[-1]                                             
                if sent[-1] == "." or sent == '':                                           
                    str1 = str1 + ""                                                        
                    break                                                                   
                else:                                                                       
                    str1 = sent + str1                                                      
                    break                                                                   
            elif len(lists[k].split(". ")) == 1 and lists[k].split(". ")[-1][-1]==".":      
                str1 = lists[k]                                                             
                break                                                                       
            else:                                                                           
                str1 = lists[k] + str1                                                      
        except:                                                                             
            pass
    
    return str1

def nextline(lists,str1,j):
    for k in range(j+1,len(lists[j])):                                                  
        try:                                                                            
            if len(lists[k].split(". ")) >1:                                            
                sent = lists[k].split(". ")[0]                                          
                str1 = str1 + sent                                                      
                break                                                                   
            elif len(lists[k].split(". ")) == 1 and lists[k].split(". ")[0][-1] == "." :
                str1 = str1 + lists[k]                                                  
                break                                                                   
            else:                                                                       
                str1 = str1 + lists[k]                                                  
        except:                                                                         
            pass
    
    return str1                                                                                                                        
    

def update_pdf(word, PASSWORD=''):
    pwd = 0
    word = word.strip()
    global count
    for filename in glob.glob(os.path.join(folder_path, '*.pdf')):
        doc = Document()
        p1 = doc.add_paragraph()
        nl = "\n"
        pdf_file = open(filename, "rb")
        pdfread = p2.PdfFileReader(pdf_file)
        if pdfread.isEncrypted:
            try:
                pdfread.decrypt(PASSWORD)
                pwd = 1
            except NotImplementedError:
                outname = filename.split["\\"][-1].split(".")[0] + "_new.pdf"
                fileout_name = filename.split("\\")[:-1] + outname
                command = f"qpdf -- password = '{PASSWORD}' --decrypt {filename} {fileout_name};"
                os.system(command)
                pdf_file = open(fileout_name, "rb")
                pdfread = p2.PdfFileReader(pdf_file)
                pwd = 1
            except:
                pdfread.decrypt('')
                pwd = 1
        for i in range(0, pdfread.getNumPages()):
            pageinfo = pdfread.getPage(i)
            listn = pageinfo.extractText().split("\n")
            for lines in listn:
                p1.add_run(lines)
                p1.add_run(nl)
            if pageinfo.extractText().casefold().find(word.casefold()) !=-1:
                lists = pageinfo.extractText().split("\n")
                for j in range(len(lists)):
                    if lists[j].casefold().count(word.casefold()) != 0:
                        
                        wf = lists[j].casefold().find(word.casefold())
                        try:
                            letr = re.findall('[A-Za-z]',lists[j][wf+len(word)])
                        except:
                            letr = ["1"]
                         
                        if len(lists[j]) == wf + len(word) or len(letr) == 0:
                            list1 = lists[j].split(". ")
                            mn = re.findall('[A-Z]',lists[j][0])
                            if len(list1) == 1 and len(mn)!= 0:
                                count = count + 1
                                excelupdate(count, filename, i, word, list1[0])
                            if len(list1) == 2 and list1[1] == "":
                                str1 = ""
                                if len(mn)!= 0:
                                    str1 = list1[0]
                                else:
                                    str1 = previous(lists, str1, j)
                                    str1 = str1 + list1[0]
                                count = count + 1
                                excelupdate(count, filename, i, word, str1)
                            if j == len(lists)-1:
                                if len(list1) == 1 and len(mn)!= 0:
                                    count = count + 1
                                    excelupdate(count, filename, i, word, lists[j])
                                        
                            else:
                                try:
                                    xy = re.findall('[A-Z]',lists[j+1][0])
                                except:
                                    xy = " "
                                if len(list1) == 1 and len(mn)!=0 and len(xy)!=0:
                                    count = count + 1
                                    excelupdate(count, filename, i, word, lists[j])
                            
                            if len(list1) == 1 and len(mn)==0 :
                                str1 = ""
                                xz = re.findall('[A-Za-z0-9]',lists[j][0])
                                if lists[j][0] == " " or len(xz) == 0:
                                    str1 = lists[j]
                                else:
                                    str1 = previous(lists, str1, j)
                                    
                                    str1 = str1 + list1[0]
                                    
                                    str1 = nextline(lists, str1, j)
                                    
                                    str1 = list1[-1] + str1
                                    pl = re.findall('[a-zA-Z0-9]', lists[j-1])
                                    if lists[j-1][-1] == "." or len(pl) == 0:
                                        str1 = lists[j]
                                count = count + 1
                                excelupdate(count, filename, i, word, str1)
                                                        
                            if len(list1)>1 and list1[1]!= "":
                                if len(list1)>2:
                                    for m in range(1,len(list1)-1):
                                        if list1[m].casefold().find(word.casefold()) != -1:
                                            count = count + 1
                                            excelupdate(count, filename, i, word, list1[m])
                                if list1[0].casefold().find(word.casefold()) != -1:
                                    str1 = ""
                                    if len(mn)!=0:
                                        str1 = list1[0]
                                    else:
                                        str1 = previous(lists, str1, j)
                                        str1 = str1 + list1[0]
                                    count = count + 1
                                    excelupdate(count, filename, i, word, str1)
                                
                                if list1[-1].casefold().find(word.casefold()) != -1:
                                    str1 = ""
                                    if list1[-1][-1] == ".":
                                        str1 = list1[-1]
                                    else:
                                        str1 = nextline(lists, str1, j)
                                        str1 = list1[-1] + str1
                                    count = count + 1
                                    excelupdate(count, filename, i, word, str1)
                        
        word_file = filename.split(".")[0] + ".docx"         #saving word file with same name as pdf                              
        doc.save(word_file)
        pdf_file.close()
    return pwd

for value in sheet.iter_rows(max_col=2, values_only=True):
    word = value[0]
    for k in sheet.iter_rows(max_col=2, values_only=True):
        password = k[1]
        if password is not None:
            pwd = update_pdf(word, password)
            if pwd == 1:
                break
        else:
            update_pdf(word)
            break

wb.save("PDFextract.xlsx") #Filename of new generated Excel File

for value in sheet.iter_rows(max_col=2, values_only=True):
    word = value[0]
    update_docx(word,folder_path)