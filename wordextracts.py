import os
import glob
import time
import re
from openpyxl import Workbook, load_workbook
from docx.enum.text import WD_COLOR_INDEX
from docx import Document

workbook = load_workbook(filename="Wkeyword.xlsx") #Enter filename of keyword
sheet = workbook.active

wb = Workbook()
ws = wb.active

ws['A1'] = 'Document Name'
ws['B1'] = 'Keyword'
ws['C1'] = 'Page No'
ws['D1'] = 'Content'


count = 0
t1 = time.time()
count1 = 1
def update_docx(word,folder_path):
    word = word.strip()
    for filename in glob.glob(os.path.join(folder_path, '*.docx')):
        count4 = 0
        count5 = 1
        global count
        global count1
        count = count +1
        doc = Document(filename)
        for i in range(0,len(doc.paragraphs)):
            for run in doc.paragraphs[i].runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    count5 = count5 + 1
                elif 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    count5 = count5 + 1
            sentence = doc.paragraphs[i].text
            sent = sentence.casefold()
            wf = sent.find(word.casefold())
            if wf != -1:
                try:
                    letr = re.findall("[a-zA-Z]", sent[wf+len(word)])
                except:
                    letr = ["1"]
                if len(sent) == wf + len(word) or len(letr) == 0:
                    count4 = count4 + 1
                    count1 = count1 + 1
                    file = filename.split("\\")
                    file = file[-1]
                    ws['A' + str(count1)] = file
                    ws['B' + str(count1)] = word
                    ws['C' + str(count1)] = count5
                    m = sentence.split(". ")
                    count2 = 0
                    for line in m:
                        p = line.casefold()
                        if p.find(word.casefold()) != -1:
                            count2 = count2+1
                            if count2 >= 2:
                                count1 = count1+1
                                ws['A' + str(count1)] = file
                                ws['B' + str(count1)] = word
                                ws['C' + str(count1)] = count5
                                ws['D' + str(count1)] = line
                            else:
                                ws['D' + str(count1)] = line
                    inline = doc.paragraphs[i].runs
                    for j in range(len(inline)):
                        if word.casefold() in inline[j].text.casefold():
                            font = inline[j].font
                            font.highlight_color = WD_COLOR_INDEX.YELLOW
                
        count_table = 0
        for table in doc.tables:
            count_table = count_table + 1
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        list1 = para.text.split(". ")
                        for line in list1:
                            wg = line.casefold().find(word.casefold())
                            if wg != -1:
                                try:
                                    letr = re.findall("[a-zA-Z]", para.text[wg+len(word)])
                                except:
                                    letr = ["1"]
                                inline = para.runs
                                for j in range(len(inline)):
                                    if word.casefold() in inline[j].text.casefold() and len(letr) == 0 or len(para.text) == wg+len(word):
                                        font = inline[j].font
                                        font.highlight_color = WD_COLOR_INDEX.YELLOW
                                        count1 = count1 + 1
                                        ws['A' + str(count1)] = file
                                        ws['B' + str(count1)] = word
                                        ws['C' + str(count1)] = "Table-" + str(count_table)
                                        ws['D' + str(count1)] = line
                                
        doc.save(filename)

for value in sheet.iter_rows(max_col=1, values_only=True):
    folder_path = 'F:\Yashraj\yes\prac' #path of word files
    word = value[0]
    update_docx(word,folder_path)

wb.save("wordextract.xlsx") #New generated excel file
t2 = time.time()
print(t2-t1)
